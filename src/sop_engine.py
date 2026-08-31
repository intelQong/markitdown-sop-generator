#!/usr/bin/env python3
"""
MarkItDown SOP Ingestion & DOCX Generator Engine
Automates multi-format document conversion and exports polished Statements of Purpose.
"""

import os
import sys
import argparse
import subprocess
from pathlib import Path

def check_dependencies():
    """Ensures all required Python packages are installed."""
    required = {
        'markitdown': 'markitdown',
        'pdfminer': 'pdfminer.six',
        'mammoth': 'mammoth',
        'docx': 'python-docx',
        'pptx': 'python-pptx',
        'openpyxl': 'openpyxl'
    }
    for mod, pkg in required.items():
        try:
            __import__(mod)
        except ImportError:
            print(f"[Setup] Installing missing dependency: {pkg}...")
            subprocess.run([sys.executable, "-m", "pip", "install", pkg], check=True)

def ingest_folder_to_markdown(folder_path: str, output_file: str = "applicant_context.md") -> str:
    """
    Converts all supported documents in a folder to a single Markdown file using MarkItDown.
    """
    from markitdown import MarkItDown
    
    md = MarkItDown()
    supported_exts = {'.pdf', '.docx', '.pptx', '.xlsx', '.csv', '.html', '.txt', '.png', '.jpg', '.jpeg', '.mp3', '.wav'}
    
    target_dir = Path(folder_path)
    if not target_dir.exists():
        raise FileNotFoundError(f"Directory not found: {folder_path}")
        
    print(f"\n[MarkItDown] Ingesting documents from: {target_dir.resolve()}...")
    
    output_lines = [
        "# APPLICANT CONTEXT & CREDENTIALS",
        f"**Source Directory:** `{target_dir.resolve()}`\n\n---\n"
    ]
    
    files_processed = 0
    for root, _, files in os.walk(target_dir):
        for f in sorted(files):
            file_p = Path(root) / f
            if file_p.suffix.lower() in supported_exts:
                print(f" -> Converting: {f} ({file_p.stat().st_size // 1024} KB)")
                output_lines.append(f"## FILE: {f}\n")
                try:
                    res = md.convert(str(file_p))
                    content = res.text_content.strip()
                    output_lines.append(content if content else "*[No text extracted / Image-only]*")
                except Exception as e:
                    output_lines.append(f"*[Extraction error: {e}]*")
                output_lines.append("\n\n---\n")
                files_processed += 1
                
    output_text = "\n".join(output_lines)
    
    out_path = Path(output_file)
    with open(out_path, "w", encoding="utf-8") as f:
        f.write(output_text)
        
    print(f"\n[MarkItDown] Successfully processed {files_processed} files into: {out_path.resolve()}\n")
    return str(out_path.resolve())

def export_sop_to_docx(sop_text: str, output_path: str, applicant_name: str = "", course_name: str = "", destination: str = ""):
    """
    Formats and exports the given SOP text into a professional Microsoft Word (.docx) document.
    """
    import docx
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    
    doc = docx.Document()
    
    # 1-inch margins
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        
    # Title
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_p.add_run("STATEMENT OF PURPOSE")
    title_run.font.name = "Calibri"
    title_run.font.size = Pt(16)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor(16, 44, 87) # Deep Navy
    title_p.paragraph_format.space_after = Pt(4)
    
    # Header Subtitle
    meta_parts = []
    if applicant_name:
        meta_parts.append(f"Applicant: {applicant_name}")
    if course_name:
        meta_parts.append(f"Target Program: {course_name}")
    if destination:
        meta_parts.append(f"Destination: {destination}")
        
    if meta_parts:
        sub_p = doc.add_paragraph()
        sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        sub_run = sub_p.add_run(" | ".join(meta_parts))
        sub_run.font.name = "Calibri"
        sub_run.font.size = Pt(10)
        sub_run.font.italic = True
        sub_run.font.color.rgb = RGBColor(100, 100, 100)
        sub_p.paragraph_format.space_after = Pt(18)
    
    # Body Paragraphs
    paragraphs = [p.strip() for p in sop_text.split("\n\n") if p.strip() and not p.strip().startswith("#") and p.strip().upper() != "STATEMENT OF PURPOSE"]
    for text in paragraphs:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.line_spacing = 1.15
        p.paragraph_format.space_after = Pt(10)
        run = p.add_run(text)
        run.font.name = "Calibri"
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(30, 30, 30)
        
    doc.save(output_path)
    print(f"\n[DOCX Generator] Formatted SOP saved to: {Path(output_path).resolve()}")

def main():
    parser = argparse.ArgumentParser(description="MarkItDown SOP Ingestion and DOCX Export Tool")
    parser.add_argument("-i", "--ingest", help="Folder path containing applicant files to convert to Markdown")
    parser.add_argument("-o", "--output-md", default="applicant_context.md", help="Output Markdown filename for ingestion")
    parser.add_argument("-e", "--export-docx", help="Text file containing SOP content to export to DOCX")
    parser.add_argument("-d", "--docx-out", default="Statement_of_Purpose.docx", help="Output .docx filename")
    parser.add_argument("--name", default="", help="Applicant full name")
    parser.add_argument("--course", default="", help="Target degree/program")
    parser.add_argument("--country", default="", help="Destination country")
    
    args = parser.parse_args()
    check_dependencies()
    
    if args.ingest:
        ingest_folder_to_markdown(args.ingest, args.output_md)
    elif args.export_docx:
        with open(args.export_docx, "r", encoding="utf-8") as f:
            content = f.read()
        export_sop_to_docx(content, args.docx_out, args.name, args.course, args.country)
    else:
        parser.print_help()

if __name__ == "__main__":
    main()